#!/usr/bin/env python3
"""
redact_word.py - Word 文档脱敏脚本
支持 .docx / .doc（含 .doc 需先转为 .docx）
依赖: python-docx
安装: pip install python-docx
"""

import sys
import re
import zipfile
import shutil
from pathlib import Path
from xml.etree import ElementTree as ET

# ---------------------------------------------------------------------------
# 脱敏规则（9类 + 银行Logo）
# ---------------------------------------------------------------------------
REPLACEMENTS = [
    (re.compile(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}'), 'XXXXX@XXXXX'),
    (re.compile(
        r'[^\x00-\xFF]{2,6}(?:省|自治区|市)?[^\x00-\xFF]{0,10}'
        r'(?:市|区)?[^\x00-\xFF]{0,10}'
        r'(?:街|路|道|巷|弄|号|大道|大街|东路|西路|南路|北路)[^\x00-\xFF]{0,30}'
    ), 'XX省XX市XX区XXXX'),
    (re.compile(r'[1-9]\d{5}(?:19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx]'),
     'XXXXXXXXXXXXXXXXXX'),
    (re.compile(r'\b(?:\d{16}|\d{17}|\d{18}|\d{19})\b'), 'XXXXXXXXXXXXXXXX'),
    (re.compile(
        r'\d{4}[-年](?:0[1-9]|1[0-2])[-月](?:0[1-9]|[12]\d|3[01])[日]?\s*'
        r'|(?:19|20)\d{2}年\d{1,2}月\d{1,2}日'
    ), 'YYYY/MM/DD'),
    (re.compile(r'\b1[3-9]\d{9}\b'), 'XXXXXXXXXXX'),
    (re.compile(r'0\d{2,3}[-\s]?\d{7,8}'), '0XX-XXXXXXXX'),
    (re.compile(
        r'(?:(?:中国|交通|招商|浦发|兴业|民生|华夏|平安|光大|广发|浙商|渤海|恒丰|'
        r'南京|宁波|杭州|深圳|上海|北京|广州|农业|建设|工商|中国)银行|'
        r'(?:农信社|信用社|农商银行|合作银行|人民银行))'
    ), '[某银行]'),
    (re.compile(r'[\u4e00-\u9fa5]{2,4}(?![a-zA-Z0-9\u4e00-\u9fa5])'), 'XXX'),
]

# Word XML 命名空间
NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
    'dc': 'http://purl.org/dc/elements/1.1/',
}


def apply_redactions(text: str) -> str:
    if not text:
        return text
    result = text
    for pattern, replacement in REPLACEMENTS:
        result = pattern.sub(replacement, result)
    return result


def redact_docx(input_path: str, output_path: str) -> None:
    """处理 .docx 文件"""
    from docx import Document

    doc = Document(input_path)
    redact_count = 0

    # 1. 处理正文所有段落
    for para in doc.paragraphs:
        for run in para.runs:
            original = run.text
            redacted = apply_redactions(original)
            if redacted != original:
                run.text = redacted
                redact_count += 1

    # 2. 处理表格单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        original = run.text
                        redacted = apply_redactions(original)
                        if redacted != original:
                            run.text = redacted
                            redact_count += 1

    # 3. 处理页眉页脚
    for section in doc.sections:
        for header in section.header.part.paragraphs:
            for run in header.runs:
                original = run.text
                redacted = apply_redactions(original)
                if redacted != original:
                    run.text = redacted
                    redact_count += 1
        if section.footer.part:
            for footer in section.footer.part.paragraphs:
                for run in footer.runs:
                    original = run.text
                    redacted = apply_redactions(original)
                    if redacted != original:
                        run.text = redacted
                        redact_count += 1

    # 4. 处理文档属性（作者、最后修改人等）
    core_props = doc.core_properties
    sensitive_props = [
        (core_props.author, 'author'),
        (core_props.last_modified_by, 'last_modified_by'),
        (getattr(core_props, 'comments', None), 'comments'),
    ]
    for val, attr_name in sensitive_props:
        if val and isinstance(val, str) and val.strip():
            redacted = apply_redactions(val)
            if redacted != val:
                setattr(core_props, attr_name, redacted)
                redact_count += 1

    # 5. 处理文本框（shapes）- 直接编辑 XML
    _redact_shapes_in_docx(input_path, output_path, redact_count)
    # 由于 python-docx 保存后可能丢失部分格式，
    # 这里直接基于 zip 层级深度编辑 XML

    # 重新用 python-docx 保存（覆盖）
    doc.save(output_path)
    print(f"[完成] 共遮盖 {redact_count} 处敏感内容，结果保存至: {output_path}")


def _redact_shapes_in_docx(input_path: str, output_path: str, base_count: int) -> int:
    """
    直接编辑 docx 内部的 XML，处理：
    - 文本框（wps:txbx / wpg:txbx）
    - 页眉页脚中的文字
    - 批注
    """
    import tempfile
    import os

    count = base_count
    tmp_dir = tempfile.mkdtemp()
    extract_dir = Path(tmp_dir) / 'docx_extracted'
    extract_dir.mkdir()

    # 解压 docx
    with zipfile.ZipFile(input_path, 'r') as z:
        z.extractall(extract_dir)

    # 处理所有 word/*.xml 文件（正文/页眉/页脚/批注/文本框）
    word_dir = extract_dir / 'word'
    if word_dir.exists():
        for xml_file in word_dir.glob('*.xml'):
            content = xml_file.read_text(encoding='utf-8')
            original = content
            for pattern, replacement in REPLACEMENTS:
                content = pattern.sub(replacement, content)
            if content != original:
                xml_file.write_text(content, encoding='utf-8')
                count += 1

    # 处理文档属性
    app_xml = extract_dir / 'docProps' / 'core.xml'
    if app_xml.exists():
        content = app_xml.read_text(encoding='utf-8')
        original = content
        for pattern, replacement in REPLACEMENTS:
            content = pattern.sub(replacement, content)
        if content != original:
            app_xml.write_text(content, encoding='utf-8')
            count += 1

    # 处理媒体文件中的银行 Logo（图片命名中含 bank/logo 等关键词）
    media_dir = word_dir / 'media'
    if media_dir.exists():
        for img_file in media_dir.iterdir():
            # 检测到图片文件名含银行相关关键词 → 替换为纯黑图
            img_name = img_file.name.lower()
            if any(k in img_name for k in ['bank', 'logo', '银行', 'logo']):
                _replace_image_with_black(img_file)
                count += 1
                print(f"  [银行Logo图片替换] {img_file.name} → 纯黑图")

    # 重新打包
    tmp_out = Path(tmp_dir) / 'output.docx'
    with zipfile.ZipFile(tmp_out, 'w', zipfile.ZIP_DEFLATED) as z:
        for file_path in extract_dir.rglob('*'):
            if file_path.is_file():
                arcname = file_path.relative_to(extract_dir)
                z.write(file_path, arcname)

    shutil.copy(tmp_out, output_path)
    shutil.rmtree(tmp_dir, ignore_errors=True)
    return count


def _replace_image_with_black(image_path: Path) -> None:
    """将图片文件替换为纯黑图（不影响文档结构）"""
    from PIL import Image
    import numpy as np

    try:
        img = Image.open(image_path)
        w, h = img.size
        black = Image.new('RGB', (max(w, 10), max(h, 10)), (0, 0, 0))
        black.save(image_path)
    except Exception as e:
        print(f"  [警告] 无法处理图片 {image_path}: {e}")


def redact_doc_to_docx(input_path: str, output_path: str) -> None:
    """将 .doc 转换为 .docx 后处理"""
    import subprocess
    stem = Path(input_path).stem
    tmp_docx = str(Path(input_path).with_name(f"{stem}_converted.docx"))
    # macOS 自带 textutil 转换 .doc → .docx
    result = subprocess.run(
        ['textutil', '-convert', 'docx', '-output', tmp_docx, input_path],
        capture_output=True, text=True
    )
    if result.returncode != 0:
        print(f"[错误] .doc 转换失败: {result.stderr}", file=sys.stderr)
        sys.exit(1)
    redact_docx(tmp_docx, output_path)
    Path(tmp_docx).unlink(missing_ok=True)


def main():
    if len(sys.argv) < 2:
        print("用法: python3 redact_word.py <输入文件路径> [输出文件路径]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    if output_file is None:
        stem = Path(input_file).stem
        output_file = str(Path(input_file).with_name(f"{stem}_脱敏.docx"))

    ext = Path(input_file).suffix.lower()
    if ext == '.doc':
        redact_doc_to_docx(input_file, output_file)
    else:
        redact_docx(input_file, output_file)


if __name__ == '__main__':
    main()
