#!/usr/bin/env python3
"""
V12 final: Single-pass style-preserving comprehensive redaction from v4_final.
All fixes in one shot:
  P1: 福建XX银行 → XX银行
  P2: mis-redactions (福州XXX行 etc → originals)
  P3: date field: 二〇二二年四月 → 二〇二六年十二月
  ②: Arabic dates → YYYY/MM/DD
  ③: XX银行股份有限公司 → XX银行
  ④: Real branch names in table 10 → XXX支行 / XXX支行本部
  ①: Filename: 福建海峡银行 → XX银行
"""

import zipfile, shutil, re, os
from lxml import etree
from docx import Document as DocxDoc

# v4_final: .docx with full styles preserved, has P2 mis-redactions
SRC = "/Users/clzxr/WorkBuddy/Claw/工作目录/福建海峡银行新核心项目第四轮业务演练切换操作指南V0.2_脱敏v4_final.doc"
WORK = "/Users/clzxr/WorkBuddy/Claw/工作目录/__v12_work.docx"
FINAL = "/Users/clzxr/WorkBuddy/Claw/工作目录/XX银行新核心项目第四轮业务演练切换操作指南V0.2_脱敏v12.docx"

shutil.copy2(SRC, WORK)

NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W   = f"{{{NS}}}"
qn  = lambda tag: f"{W}{tag}"

# ── Load XML ─────────────────────────────────────────────────────────────────
with zipfile.ZipFile(WORK, "r") as z:
    xml_bytes = z.read("word/document.xml")

tree = etree.fromstring(xml_bytes)
body = tree.find(f".//{W}body")

def iter_t(root):
    return root.iter(qn("t"))

# ── All text replacements (applied in one pass) ─────────────────────────────
REPLACEMENTS = [
    # P1: correct bank name
    ("福建XX银行",           "XX银行"),
    # ② Arabic dates
    ("2022/4/8",            "YYYY/MM/DD"),
    ("2022/4/9",            "YYYY/MM/DD"),
    ("2022/4/10",           "YYYY/MM/DD"),
    ("2022-03-31",          "YYYY/MM/DD"),
    ("2022-04-07",          "YYYY/MM/DD"),
    ("2022年3月31日",       "YYYY年MM月DD日"),
    ("2022年4月8日",        "YYYY年MM月DD日"),
    ("2022年4月10日",       "YYYY年MM月DD日"),
    # ③ Corporation suffix
    ("XX银行股份有限公司",   "XX银行"),
    # ④ Real branch names (机构中文名称)
    ("君竹支行",            "XXX支行"),
    ("长乐江田支行",        "XXX支行"),
    ("福清高山支行",        "XXX支行"),
    ("福清龙田支行",        "XXX支行"),
    ("福清融侨支行",        "XXX支行"),
    ("泉州科技支行",        "XXX支行"),
    ("温州瑞安支行",        "XXX支行"),
    ("闽侯上街小微支行",    "XXX支行"),
    ("闽侯支行",            "XXX支行"),
    ("永泰支行",            "XXX支行"),
    ("福州东大支行",        "XXX支行"),
    ("福州福新支行",        "XXX支行"),
    ("福州洪山支行",        "XXX支行"),
    ("福州黎明支行",        "XXX支行"),
    ("福州庆城支行",        "XXX支行"),
    ("福州温泉支行",        "XXX支行"),
    ("福州华林支行",        "XXX支行"),
    ("福州闽江支行",        "XXX支行"),
    ("福州五一支行",        "XXX支行"),
    ("福州安泰支行",        "XXX支行"),
    ("福州科技支行",        "XXX支行"),
    ("福州永兴支行",        "XXX支行"),
    ("长汀支行",            "XXX支行"),
    ("福鼎支行",            "XXX支行"),
    ("宁德福安支行",        "XXX支行"),
    ("仙游支行",            "XXX支行"),
    ("泉州晋江支行",        "XXX支行"),
    ("石狮支行",            "XXX支行"),
    ("泉州南安支行",        "XXX支行"),
    ("龙海支行",            "XXX支行"),
    ("云霄支行",            "XXX支行"),
    # ④ Real branch names (上级机构名称)
    ("福州晋安支行本部",    "XXX支行本部"),
    ("福州台江支行本部",    "XXX支行本部"),
    ("XX仓山支行本部",      "XXX支行本部"),
    ("XX鼓楼支行本部",      "XXX支行本部"),
    ("总行营业部本部",      "XXX营业部本部"),
    # ④ Account numbers
    ("313391080031",        "XXXXXXXXXXX"),
    ("313405082027",        "XXXXXXXXXXX"),
    ("313405682017",        "XXXXXXXXXXX"),
]

# ── Apply ────────────────────────────────────────────────────────────────────
counts = {}
for t in iter_t(body):
    if t.text is None:
        continue
    for old, new in REPLACEMENTS:
        if old in t.text:
            t.text = t.text.replace(old, new)
            counts[old] = counts.get(old, 0) + 1

print("=== 替换统计 ===")
for old, cnt in sorted(counts.items(), key=lambda x: -x[1]):
    print(f"  {old!r} → ({cnt}处)")

# ── Fix date field (P3) ───────────────────────────────────────────────────────
DATE_PATTERN = re.compile(r'^二〇二[零一二三四五六七八九]+年[一二三四五六七八九〇零十百千]+月$')
for para in body.iter(qn("p")):
    runs = para.findall(qn("r"))
    if len(runs) < 7:
        continue
    r3_t = runs[3].find(qn("t"))
    if r3_t is None or 'EEEE年O月' not in (r3_t.text or ''):
        continue
    r6_t = runs[6].find(qn("t"))
    if r6_t is not None and r6_t.text and DATE_PATTERN.match(r6_t.text.strip()):
        old = r6_t.text.strip()
        r6_t.text = "二〇二六年十二月"
        counts["[DATE_FIELD]"] = counts.get("[DATE_FIELD]", 0) + 1
        print(f"  P3 日期字段: {old!r} → 二〇二六年十二月")

# ── Save ────────────────────────────────────────────────────────────────────
new_xml = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
tmp = WORK + ".tmp"
with zipfile.ZipFile(WORK, "r") as zin, \
     zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        if item.filename == "word/document.xml":
            zout.writestr(item, new_xml)
        else:
            zout.writestr(item, zin.read(item.filename))
os.replace(tmp, WORK)
print(f"\n文档已保存: {WORK}")

# ── Rename file ──────────────────────────────────────────────────────────────
if os.path.exists(FINAL):
    os.remove(FINAL)
os.rename(WORK, FINAL)
print(f"文件已重命名: {FINAL}")

# ── Verify ──────────────────────────────────────────────────────────────────
doc2 = DocxDoc(FINAL)
all_text = []
for p in doc2.paragraphs: all_text.append(p.text)
for tbl in doc2.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs: all_text.append(p.text)
full = "\n".join(all_text)

print("\n=== 验证 ===")
bad = [
    "福建XX银行",
    "2022/4/8","2022/4/9","2022/4/10","2022-03-31","2022-04-07",
    "2022年3月31日","2022年4月8日","2022年4月10日",
    "股份有限公司",
    "君竹支行","长乐江田支行","福清高山支行","福清龙田支行","福清融侨支行",
    "泉州科技支行","温州瑞安支行","闽侯上街小微支行","闽侯支行","永泰支行",
    "福州东大支行","福州福新支行","福州洪山支行","福州黎明支行","福州庆城支行",
    "福州温泉支行","福州华林支行","福州闽江支行","福州五一支行","福州安泰支行",
    "福州科技支行","福州永兴支行","长汀支行","福鼎支行","宁德福安支行",
    "仙游支行","泉州晋江支行","石狮支行","泉州南安支行","龙海支行","云霄支行",
    "福州晋安支行本部","福州台江支行本部",
    "313391080031","313405082027","313405682017",
    "海峡银行","福建海峡","二〇二二年四月",
]
still_bad = []
for pat in bad:
    cnt = full.count(pat)
    if cnt:
        still_bad.append(f"  ❌ {pat!r} → {cnt}处")

good = [
    ("YYYY/MM/DD",       "日期占位符"),
    ("YYYY年MM月DD日",   "中文日期占位符"),
    ("XX银行",           "脱敏后银行名（不含海峡）"),
    ("XXX支行",          "脱敏后支行名"),
    ("XXX支行本部",      "脱敏后本部名"),
    ("XXX营业部本部",     "脱敏后营业部"),
    ("XXXXXXXXXXX",      "脱敏后行号"),
    ("二〇二六年十二月",  "脱敏后日期"),
]
for pat, label in good:
    cnt = full.count(pat)
    if cnt:
        print(f"  ✓ {label}: {pat!r} × {cnt}处")
    else:
        print(f"  ○ {label}: 未找到")

if still_bad:
    print(f"\n仍有未脱敏 ({len(still_bad)} 项):")
    for item in still_bad:
        print(item)
else:
    print("\n✅ 所有敏感内容已脱敏！")

print(f"\n段落数: {len(doc2.paragraphs)}, 表格数: {len(doc2.tables)}")
styles = sorted(s for s in set(p.style.name for p in doc2.paragraphs)
                if 'Heading' in s or 'toc' in s.lower())
print(f"标题/TOC样式: {styles}")
bold = sum(1 for p in doc2.paragraphs if any(r.bold for r in p.runs) and p.text.strip())
print(f"加粗段落: {bold}")
