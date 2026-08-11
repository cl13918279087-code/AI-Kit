#!/usr/bin/env python3
"""
V11: Style-preserving redaction via direct XML manipulation.
Preserves ALL Word formatting: heading styles, bold, font size, table styles, etc.

Fixes:
  P1: 福建XX银行 → XX银行
  P2: N/A (v4_final already clean of those patterns)
  P3: 日期字段（二〇二二年四月 → 二〇二六年十二月）
"""

import zipfile, shutil, re, os
from lxml import etree

SRC = "/Users/clzxr/WorkBuddy/Claw/工作目录/福建海峡银行新核心项目第四轮业务演练切换操作指南V0.2_脱敏v4_final.doc"
DST = "/Users/clzxr/WorkBuddy/Claw/工作目录/福建海峡银行新核心项目第四轮业务演练切换操作指南V0.2_脱敏v11.docx"

shutil.copy2(SRC, DST)

NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W  = f"{{{NS}}}"
qn = lambda tag: f"{W}{tag}"

# ── Load ──────────────────────────────────────────────────────────────────────
with zipfile.ZipFile(DST, "r") as z:
    xml_bytes = z.read("word/document.xml")

tree = etree.fromstring(xml_bytes)
body = tree.find(f".//{W}body")   # OOXML uses default ns

# ── Text-replacement helpers ───────────────────────────────────────────────────

def iter_t_elements(root):
    """Yield all <w:t> elements in document order."""
    return root.iter(qn("t"))

def replace_in_t(t_elem, old, new):
    """Replace old→new in a <w:t> element. Returns count."""
    if t_elem.text is None:
        return 0
    if old not in t_elem.text:
        return 0
    t_elem.text = t_elem.text.replace(old, new)
    return 1

# ── P1 fixes ──────────────────────────────────────────────────────────────────
P1 = [("福建XX银行", "XX银行")]

# ── P3: Date field ─────────────────────────────────────────────────────────────
# Run structure in v4_final:
#   run[0]: <w:t>' 编制日期：'
#   run[1]: <w:fldChar fldCharType="begin"/>
#   run[2]: <w:instrText>' '
#   run[3]: <w:t>'TIME \@ "EEEE年O月"'
#   run[4]: <w:instrText>' '
#   run[5]: <w:fldChar fldCharType="separate"/>
#   run[6]: <w:t>'二〇二二年四月'   ← cached result to mask
#   run[7]: <w:fldChar fldCharType="end"/>

DATE_PATTERN = re.compile(r'^二〇二[零一二三四五六七八九]+年[一二三四五六七八九〇零十百千]+月$')

# ── Pass 1: text replacement ───────────────────────────────────────────────────
changes_p1 = 0
for t in iter_t_elements(body):
    for old, new in P1:
        changes_p1 += replace_in_t(t, old, new)

print(f"P1 (福建XX银行→XX银行): {changes_p1} <w:t> nodes changed")

# ── Pass 2: date field fix ─────────────────────────────────────────────────────
# Walk paragraph by paragraph to find the date field paragraph
changes_p3 = 0
for para in body.iter(qn("p")):
    runs = para.findall(qn("r"))
    if len(runs) < 7:
        continue

    # Check if run[3] contains the date format code
    r3_t = runs[3].find(qn("t"))
    if r3_t is None or 'EEEE年O月' not in (r3_t.text or ''):
        continue

    # Check if run[6] has the date result
    r6_t = runs[6].find(qn("t"))
    if r6_t is not None and r6_t.text and DATE_PATTERN.match(r6_t.text.strip()):
        old_date = r6_t.text.strip()
        r6_t.text = "二〇二六年十二月"
        changes_p3 += 1
        print(f"  P3: {old_date!r} → 二〇二六年十二月")

print(f"P3 (日期脱敏): {changes_p3} fields changed")

# ── Save ──────────────────────────────────────────────────────────────────────
new_xml = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)

tmp = DST + ".tmp"
with zipfile.ZipFile(DST, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        if item.filename == "word/document.xml":
            zout.writestr(item, new_xml)
        else:
            zout.writestr(item, zin.read(item.filename))

os.replace(tmp, DST)
print(f"\nSaved: {DST}")

# ── Verify ─────────────────────────────────────────────────────────────────────
from docx import Document as DocxDoc

doc2 = DocxDoc(DST)
all_text = []
for p in doc2.paragraphs:
    all_text.append(p.text)
for tbl in doc2.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                all_text.append(p.text)

full = "\n".join(all_text)

print("\n=== 验证 ===")
bad = ["福建XX银行","福州XXX行","宁XXX行","泉XXX行","漳XXX行","龙XXX行","长XXX行","上XXX行","网XXX行","XXX行ATM","新核心系统XXX项目"]
for pat in bad:
    cnt = full.count(pat)
    if cnt:
        print(f"  ❌ STILL BAD: {pat!r} → {cnt}处")

good_items = [
    ("XX银行",              "XX银行替换"),
    ("福州分行",            "福州分行"),
    ("宁德分行",            "宁德分行"),
    ("泉州分行",            "泉州分行"),
    ("漳州分行",            "漳州分行"),
    ("龙岩分行",            "龙岩分行"),
    ("长乐分行",            "长乐分行"),
    ("网上银行",            "网上银行"),
    ("新核心系统建设项目",  "新核心系统建设项目"),
    ("二〇二六年十二月",    "脱敏后日期"),
]
for pat, label in good_items:
    cnt = full.count(pat)
    if cnt:
        print(f"  ✓  {label}: {pat!r} → {cnt}处")
    else:
        print(f"  ○  {label}: {pat!r} → 0处（未出现）")

print("\n=== 格式验证 ===")
# Check that heading styles are preserved
styles_used = set(p.style.name for p in doc2.paragraphs)
heading_styles = [s for s in styles_used if 'Heading' in s or 'toc' in s.lower()]
print(f"  标题/TOC样式: {sorted(heading_styles)}")
bold_count = sum(1 for p in doc2.paragraphs if any(r.bold for r in p.runs) and p.text.strip())
print(f"  加粗段落数: {bold_count}")
print(f"  总段落数: {len(doc2.paragraphs)}, 表格数: {len(doc2.tables)}")
