#!/usr/bin/env python3
"""
Post-process v12: fix remaining corporation suffix and
handle dates split across multiple XML runs.
"""

import zipfile, re, os, shutil
from lxml import etree
from docx import Document as DocxDoc

DST = "/Users/clzxr/WorkBuddy/Claw/工作目录/XX银行新核心项目第四轮业务演练切换操作指南V0.2_脱敏v12.docx"
WORK = "/Users/clzxr/WorkBuddy/Claw/工作目录/__v12b_work.docx"

shutil.copy2(DST, WORK)

NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W   = f"{{{NS}}}"
qn  = lambda tag: f"{W}{tag}"

with zipfile.ZipFile(WORK, "r") as z:
    xml_bytes = z.read("word/document.xml")

tree = etree.fromstring(xml_bytes)
body = tree.find(f".//{W}body")

# ── Fix 1: "股份有限公司" → "XXXX" ──────────────────────────────────────────
# v4_final already removed "XX银行" prefix, leaving "股份有限公司XX鼓楼支行"
# → should become "XXXXXX鼓楼支行"
CORP_REPLACEMENTS = [
    ("股份有限公司XX鼓楼支行", "XXXXXX鼓楼支行"),
    ("股份有限公司XX新罗支行", "XXXXXX新罗支行"),
    ("股份有限公司漳平支行",   "XXXXXX漳平支行"),
    ("股份有限公司XX金山支行", "XXXXXX金山支行"),
    ("股份有限公司XX仓山支行", "XXXXXX仓山支行"),
    ("股份有限公司XX杨桥支行", "XXXXXX杨桥支行"),
    ("股份有限公司漳浦支行",   "XXXXXX漳浦支行"),
]

corp_fixed = 0
for t in body.iter(qn("t")):
    if t.text is None:
        continue
    for old, new in CORP_REPLACEMENTS:
        if old in t.text:
            t.text = t.text.replace(old, new)
            corp_fixed += 1
print(f"Corporation suffix fixed: {corp_fixed} nodes")

# ── Fix 2: Handle dates split across XML runs ───────────────────────────────
# Collect all <w:p> paragraphs, merge run texts, apply replacements, split back
DATE_PATTERN = re.compile(r'2022[/-]0?(\d+)[/-]0?(\d+)|2022年(\d+)月(\d+)日')

def merge_runs(para):
    """Merge all run texts in a paragraph into one string, remembering positions."""
    runs = para.findall(qn("r"))
    parts = []
    run_map = []  # (run_idx, char_start, char_end)
    pos = 0
    for ri, r in enumerate(runs):
        t = r.find(qn("t"))
        if t is not None and t.text:
            start = pos
            parts.append(t.text)
            pos += len(t.text)
            run_map.append((ri, start, pos))
        else:
            # Empty run, skip
            pass
    return ''.join(parts), runs, run_map

def split_back(para, new_text, runs, run_map):
    """Split new_text back across runs, clearing extra runs."""
    if not run_map:
        return
    # Put as much as possible into first run
    runs[run_map[0][0]].find(qn("t")).text = new_text
    for ri, _, _ in run_map[1:]:
        t = runs[ri].find(qn("t"))
        if t is not None:
            t.text = ''

DATE_FIXES = {
    "2022/4/8":  "YYYY/MM/DD",
    "2022/4/9":  "YYYY/MM/DD",
    "2022/4/10": "YYYY/MM/DD",
    "2022-03-31": "YYYY/MM/DD",
    "2022-04-07": "YYYY/MM/DD",
    "2022年3月31日": "YYYY年MM月DD日",
    "2022年4月8日":  "YYYY年MM月DD日",
    "2022年4月10日": "YYYY年MM月DD日",
}

dates_fixed = 0
for para in body.iter(qn("p")):
    runs = para.findall(qn("r"))
    if not runs:
        continue

    # Collect all text from runs in this paragraph
    texts = []
    for r in runs:
        t = r.find(qn("t"))
        texts.append(t.text if t is not None and t.text else "")

    full = "".join(texts)

    # Check if any date is split across runs
    needs_fix = any(d in full for d in DATE_FIXES)
    if not needs_fix:
        continue

    # Apply all date fixes to merged text
    modified = full
    for old, new in DATE_FIXES.items():
        if old in modified:
            modified = modified.replace(old, new)
            dates_fixed += 1

    if modified == full:
        continue

    # Rebuild runs: put all text in first non-None run, clear others
    # Find first run with a <w:t> element
    first_run_idx = None
    for ri, r in enumerate(runs):
        t = r.find(qn("t"))
        if t is not None:
            first_run_idx = ri
            break

    if first_run_idx is not None:
        t = runs[first_run_idx].find(qn("t"))
        t.text = modified
        # Clear remaining runs
        for ri in range(first_run_idx + 1, len(runs)):
            t2 = runs[ri].find(qn("t"))
            if t2 is not None:
                t2.text = ''

print(f"Dates fixed (run-split): {dates_fixed}")

# ── Save ────────────────────────────────────────────────────────────────────
new_xml = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
tmp = WORK + ".tmp"
with zipfile.ZipFile(WORK, "r") as zin, \
     zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        if item.filename == "word/document.xml":
            zout.writestr(item, new_xml)
        else:
            zout.writestr(item, zin.read(item.filename))
os.replace(tmp, WORK)
print(f"Saved: {WORK}")

# Move to final
if os.path.exists(DST):
    os.remove(DST)
os.rename(WORK, DST)
print(f"Final: {DST}")

# ── Verify ──────────────────────────────────────────────────────────────────
import shutil
doc2 = DocxDoc(DST)
all_text = []
for p in doc2.paragraphs: all_text.append(p.text)
for tbl in doc2.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs: all_text.append(p.text)
full = "\n".join(all_text)

print("\n=== 验证 ===")
bad = [
    "2022/4/8","2022/4/9","2022/4/10","2022-03-31","2022-04-07",
    "2022年3月31日","2022年4月8日","2022年4月10日",
    "股份有限公司",
    "海峡银行","福建海峡","二〇二二年四月",
]
still_bad = []
for pat in bad:
    cnt = full.count(pat)
    if cnt:
        still_bad.append(f"  ❌ {pat!r} → {cnt}处")

good = [
    ("YYYY/MM/DD",       "日期占位符"),
    ("YYYY年MM月DD日",   "中文日期占位符"),
    ("XX银行",           "脱敏后银行名"),
    ("XXXXXX鼓楼支行",   "XXXXXX鼓楼支行"),
]
for pat, label in good:
    cnt = full.count(pat)
    if cnt:
        print(f"  ✓ {label}: {pat!r} × {cnt}处")

if still_bad:
    print(f"\n仍有 ({len(still_bad)} 项):")
    for item in still_bad: print(item)
else:
    print("\n✅ 全部脱敏完成！")

print(f"\n段落: {len(doc2.paragraphs)}, 表格: {len(doc2.tables)}, 样式: {len(set(p.style.name for p in doc2.paragraphs))}")
